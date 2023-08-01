import csv
import os
import shutil
from dotenv import load_dotenv
from docx import Document
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from dateutil.relativedelta import relativedelta

from classes.story import Story, StoryGroup
import classes.globals as g


class ReportWriter(object):
    def __init__(self):
        self.get_date_string()
        self.get_config()
        self.get_latest_csv()
        self.read_csv()
        self.group_stories()

    def get_date_string(self):
        d = datetime.now()
        day = int(d.strftime("%d"))
        if day < 15:
            d = d - relativedelta(months=1)
        self.year = d.strftime("%y")
        self.month = d.strftime("%m")
        self.month_name = d.strftime("%b")

    def get_config(self):
        load_dotenv('.env')
        try:
            self.write_story_points = int(os.getenv('write_story_points'))
        except Exception as e:
            print(e.args)
            self.write_story_points = 0

        self.resources_folder = os.path.join(os.getcwd(), "resources")
        self.jira_folder = os.path.join(self.resources_folder, "jira")
        self.report_folder = os.path.join(self.resources_folder, "report")
        self.template_folder = os.path.join(self.resources_folder, "template")

        # Make folders
        if not os.path.isdir(self.resources_folder):
            os.mkdir(self.resources_folder)
        if not os.path.isdir(self.jira_folder):
            os.mkdir(self.jira_folder)
        if not os.path.isdir(self.report_folder):
            os.mkdir(self.report_folder)

        self.stories = []

        self.template_filename = os.path.join(self.template_folder, "report_template.docx")
        self.report_filename = os.path.join(self.report_folder, "OTT Monthly Deliverables {year}-{month}.docx".format(
            year=self.year,
            month=self.month,
        ))
        self.title = "OTT Monthly Deliverables {month} {year}".format(
            year=self.year,
            month=self.month_name,
        )

    def read_csv(self):
        with open(self.csv_file, encoding="utf8") as csv_file:
            csv_reader = csv.reader(csv_file, delimiter=',')
            line_count = 0
            for row in csv_reader:
                if line_count == 0:
                    self.get_keys(row)
                    line_count += 1
                else:
                    story = Story(row)
                    self.stories.append(story)
                    line_count += 1
            print(f'Processed {line_count} lines.')

    def get_keys(self, row):
        for item in g.fields:
            field = g.fields[item]
            column_index = -1
            found = False
            for cell in row:
                column_index += 1
                if cell == field["header"]:
                    field["actual"] = column_index
                    found = True
                    break

            if not found:
                field["actual"] = field["default"]

    def group_stories(self):
        self.story_groups = []
        previous_epic = "dummy"
        story_group = StoryGroup()

        if (len(self.stories)) > 0:
            for story in self.stories:
                if story.epic != previous_epic:
                    if len(story_group.stories) > 0:
                        self.story_groups.append(story_group)
                    story_group = StoryGroup()
                    story_group.epic = story.epic
                    story_group.stories.append(story)
                else:
                    story_group.stories.append(story)

                previous_epic = story.epic

        story_group.epic = story.epic
        self.story_groups.append(story_group)

    def set_repeat_table_header(self, row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def color_row(self, row, rgb="f0f0f0"):
        for cell in row.cells:
            shading_elm_2 = parse_xml(r'<w:shd {} w:fill="f0f0f0"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm_2)

    def write(self):
        self.document = Document(self.template_filename)
        self.document.add_heading(self.title, 0)

        for story_group in self.story_groups:
            total = 0
            self.document.add_heading(story_group.epic, 1)

            if self.write_story_points:
                widths = (Cm(2.5), Cm(12), Cm(1.5))
                headers = ["Story", "Description", "Points"]
                column_count = 3
                row_increment = 2
            else:
                widths = (Cm(2.5), Cm(13.5))
                headers = ["Story", "Description"]
                column_count = 2
                row_increment = 1

            # Add the table
            table = self.document.add_table(rows=len(story_group.stories) + row_increment, cols=column_count)
            table.style = "List Table 3"

            # Set table widths
            for row in table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

            # Set table headers
            self.set_repeat_table_header(table.rows[0])
            hdr_cells = table.rows[0].cells
            for i in range(0, len(headers)):
                p = hdr_cells[i].paragraphs[0]
                if i == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.text = headers[i]

                # hdr_cells[i].text = headers[i]

            # Fill table cells
            row_count = 1
            for story in story_group.stories:
                cells = table.rows[row_count].cells
                cells[0].text = story.key
                cells[1].text = story.summary
                if self.write_story_points:
                    p = cells[2].paragraphs[0]
                    if story.story_points != "":
                        total += int(story.story_points)
                        p.text = story.story_points
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_count += 1

            # Write the totals row
            if self.write_story_points:
                row = table.rows[row_count]
                self.color_row(row, "f0f0f0")
                cells = row.cells
                cells[0].text = "TOTAL"
                p = cells[2].paragraphs[0]
                runner = p.add_run(str(total))
                runner.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.save(self.report_filename)
        self.copy_to_governance_folder()

    def copy_to_governance_folder(self):
        load_dotenv('.env')
        self.governance_folder = os.getenv('governance_folder')
        if self.governance_folder is not None and self.governance_folder != "":
            if os.path.isdir(self.governance_folder):
                shutil.copy(self.report_filename, self.governance_folder)

    def get_latest_csv(self):
        files = os.listdir(self.jira_folder)
        paths = [os.path.join(self.jira_folder, basename) for basename in files if basename.endswith(".csv")]
        self.csv_file = max(paths, key=os.path.getctime)
